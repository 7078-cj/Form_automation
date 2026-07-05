from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def replace_text(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            placeholder = f"({key})"
            if placeholder in paragraph.text:
                # Replace while preserving formatting as much as possible
                for run in paragraph.runs:
                    run.text = run.text.replace(placeholder, str(value))

    # Replace text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f"({key})"
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(placeholder, str(value))

def insert_participants_table(doc: Document, participants):
    """
    Replace the paragraph containing (participants_table)
    with a dynamically generated table.
    """

    for i, paragraph in enumerate(doc.paragraphs):
        if "(participants_table)" not in paragraph.text:
            continue

        # Remove placeholder
        paragraph.text = ""

        cols = len(participants)
        rows = max(len(team["members"]) for team in participants) + 1

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # Header
        for col, team in enumerate(participants):
            cell = table.cell(0, col)
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = p.add_run(team["team_name"])
            run.bold = True
            run.font.size = Pt(11)

        # Members
        for col, team in enumerate(participants):
            for row, member in enumerate(team["members"], start=1):
                cell = table.cell(row, col)
                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run(member)

        # Move the table where the placeholder was
        tbl = table._tbl
        paragraph._p.addnext(tbl)

        break